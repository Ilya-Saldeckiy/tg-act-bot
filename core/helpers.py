import os, datetime
from pathlib import Path
from PIL import Image
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import subprocess


def load_env_file(file_path: str):
    if not file_path.is_file():
        print(f"Файл {file_path} не найден.")
        return

    with open(file_path, "r") as env_file:
        for line in env_file:
            line = line.strip()
            if not line or line.startswith("#"):
                continue

            key, value = line.split("=", 1)
            os.environ[key] = value


def convert_docx_to_pdf(input_file, output_dir):
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, input_file])
    output_file = os.path.join(output_dir, os.path.splitext(os.path.basename(input_file))[0] + ".pdf")

    return output_file


def add_paragraph(
    doc, text: str, font_size: int = 13, bold: bool = False, color: RGBColor = RGBColor(0, 0, 0), space_after: int = 15, alignment=None
):
    """
    Добавляет абзац в документ с заданными параметрами.
    :param doc: Документ, в который добавляется абзац.
    :param text: Текст абзаца.
    :param font_size: Размер шрифта.
    :param bold: Жирный шрифт.
    :param color: Цвет текста.
    :param space_after: Отступ после абзаца (в пунктах).
    :param alignment: Выравнивание абзаца.
    :return: Созданный абзац.
    """
    paragraph = doc.add_paragraph()
    if alignment:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def create_docx_file(user_id: int, db: Session = None):
    from core.models import ItemDB, Users

    ACTS_DIR = Path("acts")
    ACTS_DIR.mkdir(exist_ok=True)

    data_obj = db.query(ItemDB).filter(ItemDB.tg_id == user_id).all()

    if data_obj:
        data = data_obj[-1]

    # Создание документа
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"  # Устанавливаем шрифт
    style.font.size = Pt(13)  # Устанавливаем размер шрифта
    style.font.color.rgb = RGBColor(0, 0, 0)  # Устанавливаем цвет текста (чёрный)

    current_datetime = datetime.datetime.now()
    current_date = current_datetime.strftime("%d/%m/%Y")
    current_time = current_datetime.strftime("%H:%M")

    current_user = db.query(Users).filter(Users.tg_id == user_id).first()    
    current_user_full_name = current_user.full_name if current_user else "Тестовый тест"

    if current_user_full_name:
        name_for_act = "".join(word[0] for word in current_user_full_name.split()[:2])

    title = f"АКТ №{current_datetime.strftime('%d%m')}{name_for_act if name_for_act else "ХХ"}{data.id if data.id >= 10 else '0' + str(data.id)} {data.title if data.title else 'НАЗВАНИЕ АКТА'}"

    # Начинаем наполнение документа
    add_paragraph(doc, f"Объект: {data.object_name}", bold=True, space_after=0)
    add_paragraph(doc, f"Раздел проекта: {data.project_name}", bold=True, space_after=0)
    add_paragraph(doc, f"Компания: {data.company_name}", bold=True, space_after=15)

    # Добавляем дату и имя составителя
    add_paragraph(
        doc,
        f"Дата составления: {current_date} время: {current_time}\nСоставил инженер СК: ____________________ {current_user_full_name}",
        color=RGBColor(158, 158, 158),
        space_after=30,
    )

    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.left_indent = Inches(0.7)
    heading.paragraph_format.right_indent = Inches(0.7)

    heading_run = heading.add_run("АКТ выявленных недостатков, дефектов и несоответствий работ ")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    heading_run = heading.add_run(
        f"№{current_datetime.strftime('%d%m')}{name_for_act if name_for_act else 'ХХ'}{data.id if data.id >= 10 else '0' + str(data.id)}"
    )
    heading_run.font.size = Pt(16)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    heading_run.font.bold = True
    heading_run.underline = True

    # Добавляем подзаголовок
    add_paragraph(doc, "В результате проведенной проверки установлены следующие нарушения:", space_after=10)

    # Добавление данных из объекта
    for item_id, item in data.data_obj.items():
        counter = item_id

        for text in item.get("texts"):
            paragraph = doc.add_paragraph(f"{counter}. {text}")
            paragraph.paragraph_format.space_after = Pt(15)

        if not item.get("texts"):
            paragraph = doc.add_paragraph(f"{counter}. ")
            paragraph.paragraph_format.space_after = Pt(15)

        # Добавление изображений
        for photo in item.get("photos"):
            if Path(photo).exists():
                try:
                    with Image.open(photo) as image:
                        width_px, height_px = image.size

                        # Стандартное разрешение в DPI (точки на дюйм)
                        dpi = 96

                        width_inch = width_px / dpi
                        height_inch = height_px / dpi

                        # Ограничиваем размеры изображения
                        max_width_inch = 4
                        max_height_inch = 5.0

                        if width_inch > max_width_inch:
                            scale = max_width_inch / width_inch
                            width_inch *= scale
                            height_inch *= scale

                        if height_inch > max_height_inch:
                            scale = max_height_inch / height_inch
                            width_inch *= scale
                            height_inch *= scale

                        # Добавляем изображение в документ
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(photo, width=Inches(width_inch), height=Inches(height_inch))

                        # Выравниваем абзац по центру
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Устанавливаем отступ после абзаца в 15 pt
                        paragraph.paragraph_format.space_after = Pt(15)
                except Exception as e:
                    paragraph = doc.add_paragraph(f"Ошибка при добавлении фото: {str(e)}")
                    paragraph.paragraph_format.space_after = Pt(15)

    # Сохранение документа
    filename = title + ".docx"
    file_path = ACTS_DIR / filename

    doc.save(file_path)

    # Сохранение пути в базе данных
    if os.path.exists(file_path):

        file_path_pdf = convert_docx_to_pdf(str(file_path), "acts/")
        db.query(ItemDB).filter(ItemDB.id == data.id).update({"file_path": str(file_path), "file_path_pdf": str(file_path_pdf)})
        db.commit()

        return title, str(file_path), str(file_path_pdf)
    else:
        return None
